"""
Exporters для документов в различные форматы
"""

from typing import Optional
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

from app.models import Document


class DOCXExporter:
    """Экспорт документов в DOCX формат"""

    def __init__(self):
        pass

    def export_document(self, document: Document) -> BytesIO:
        """
        Экспорт документа в DOCX

        :param document: Документ для экспорта
        :return: BytesIO с DOCX содержимым
        """
        # Создать новый DOCX документ
        docx = DocxDocument()

        # Заголовок документа
        title = docx.add_heading(document.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Метаданные
        if document.template_id:
            p = docx.add_paragraph()
            p.add_run("Template: ").bold = True
            p.add_run(document.template_id)

        p = docx.add_paragraph()
        p.add_run("Status: ").bold = True
        p.add_run(document.status)

        p = docx.add_paragraph()
        p.add_run("Created: ").bold = True
        p.add_run(document.created_at.strftime("%Y-%m-%d %H:%M"))

        # Разделитель
        docx.add_paragraph("_" * 60)
        docx.add_page_break()

        # Добавить блоки
        for doc_block in sorted(document.blocks, key=lambda b: b.order):
            # Заголовок блока по уровню
            if doc_block.level == 1:
                heading = docx.add_heading(f"Block: {doc_block.block_id}", level=1)
            elif doc_block.level == 2:
                heading = docx.add_heading(f"Block: {doc_block.block_id}", level=2)
            else:
                heading = docx.add_heading(f"Block: {doc_block.block_id}", level=3)

            # Содержимое блока
            content_para = docx.add_paragraph(doc_block.content)
            content_para.paragraph_format.left_indent = Inches(0.25 * doc_block.level)
            content_para.paragraph_format.space_after = Pt(12)

        # Футер с контекстом (если есть)
        if document.context:
            docx.add_page_break()
            docx.add_heading("Context Variables", level=1)
            for key, value in document.context.items():
                p = docx.add_paragraph()
                p.add_run(f"{key}: ").bold = True
                p.add_run(str(value))

        # Сохранить в BytesIO
        buffer = BytesIO()
        docx.save(buffer)
        buffer.seek(0)

        return buffer


# Singleton instance
docx_exporter = DOCXExporter()
